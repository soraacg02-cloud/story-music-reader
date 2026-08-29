import io
import re
from datetime import datetime
from urllib.parse import quote, unquote
import cloudinary
import cloudinary.api
import cloudinary.uploader
import docx
import requests
import streamlit as st

# 1. 頁面基本配置
st.set_page_config(
    page_title="雲端沉浸式故事音樂書櫃", page_icon="📚", layout="wide"
)

# 2. Session State 記憶狀態初始化
if "selected_book_id" not in st.session_state:
    st.session_state.selected_book_id = None
if "ch_index" not in st.session_state:
    st.session_state.ch_index = 0
if "theme_mode" not in st.session_state:
    st.session_state.theme_mode = "🌙 黑底模式"
if "auto_play" not in st.session_state:
    st.session_state.auto_play = True
if "loop_play" not in st.session_state:
    st.session_state.loop_play = True
if "max_chapters" not in st.session_state:
    st.session_state.max_chapters = 1
if "chapter_titles" not in st.session_state:
    st.session_state.chapter_titles = []
if "reading_pct" not in st.session_state:
    st.session_state.reading_pct = 0

# ---------------- 【左側欄順序 1】：視覺風格 ----------------
st.sidebar.header("🎨 視覺風格")
st.session_state.theme_mode = st.sidebar.radio(
    "選擇閱讀配色：",
    ["🌙 黑底模式", "☀️ 白底模式"],
    index=0 if st.session_state.theme_mode == "🌙 黑底模式" else 1,
    key="theme_radio",
)

# 3. 高對比度與主題顏色 CSS 設定
if st.session_state.theme_mode == "🌙 黑底模式":
    bg_col, sidebar_bg, card_bg, border_col, text_col, button_bg = (
        "#0e1117",
        "#161920",
        "#1f232a",
        "#30363d",
        "#e0e0e0",
        "#262c36",
    )
else:
    bg_col, sidebar_bg, card_bg, border_col, text_col, button_bg = (
        "#f9f9fb",
        "#ffffff",
        "#ffffff",
        "#e1e4e8",
        "#1f232a",
        "#ffffff",
    )

css_style = f"""
<style>
    /* 全局背景與側邊欄背景 */
    .stApp {{ background-color: {bg_col} !important; color: {text_col}; }}
    [data-testid="stSidebar"], [data-testid="stSidebarContent"] {{ background-color: {sidebar_bg} !important; }}
    
    /* 側邊欄文字 */
    [data-testid="stSidebar"] label, [data-testid="stSidebar"] p, [data-testid="stSidebar"] span {{ color: {text_col} !important; }}

    /* 卡片與容器樣式 */
    div[data-testid="stContainer"] {{ 
        background-color: {card_bg} !important; 
        border: 1px solid {border_col} !important; 
        border-radius: 12px; 
    }}

    /* 按鈕樣式限制 */
    div.stButton > button {{ 
        background-color: {button_bg} !important; 
        color: {text_col} !important; 
        border: 1px solid {border_col} !important; 
        font-weight: bold !important; 
        border-radius: 8px !important; 
    }}

    /* Popover 內彈窗專屬樣式修正 */
    div[data-testid="stPopoverBody"] {{
        background-color: {card_bg} !important;
        border: 1px solid {border_col} !important;
        border-radius: 12px !important;
    }}
    div[data-testid="stPopoverBody"] * {{ 
        color: {text_col} !important; 
    }}

    /* 文章內文專屬樣式：控制句子/段落間距與行高 */
    .story-paragraph {{
        color: {text_col} !important;
        font-size: 1.15rem !important;
        line-height: 1.95 !important;
        margin-bottom: 1.25rem !important;
        white-space: pre-wrap !important;
        word-break: break-word !important;
    }}
</style>
"""
st.markdown(css_style, unsafe_allow_html=True)


# 4. 初始化 Cloudinary 連線
def init_cloudinary():
    if "cloudinary" not in st.secrets:
        return False
    cfg = st.secrets["cloudinary"]
    cloudinary.config(
        cloud_name=str(cfg.get("cloud_name", "")).strip(),
        api_key=str(cfg.get("api_key", "")).strip(),
        api_secret=str(cfg.get("api_secret", "")).strip(),
        secure=True,
    )
    return True


# 檔案大小可讀化轉換函式
def format_file_size(size_in_bytes):
    if not size_in_bytes or size_in_bytes <= 0:
        return "大小未知"
    if size_in_bytes < 1024:
        return f"{size_in_bytes} Bytes"
    elif size_in_bytes < 1024 * 1024:
        return f"{size_in_bytes / 1024:.1f} KB"
    else:
        return f"{size_in_bytes / (1024 * 1024):.2f} MB"


# 5. 安全狀態同步回呼函式
def prev_chapter_cb():
    if st.session_state.ch_index > 0:
        st.session_state.ch_index -= 1
        st.session_state.reading_pct = 0


def next_chapter_cb():
    if st.session_state.ch_index < st.session_state.max_chapters - 1:
        st.session_state.ch_index += 1
        st.session_state.reading_pct = 0


def on_radio_change_cb():
    # 同時相容頂部與側邊欄的選單 key
    selected = st.session_state.get("top_popover_radio") or st.session_state.get("sb_popover_radio")
    titles = st.session_state.get("chapter_titles", [])
    if selected in titles:
        st.session_state.ch_index = titles.index(selected)
        st.session_state.reading_pct = 0


def on_slider_change_cb():
    st.session_state.reading_pct = st.session_state.get("sb_slider_pct", 0)


# 6. 快取網路請求與 Word 文件解析
@st.cache_data(ttl=60)
def fetch_cloudinary_catalog():
    raw_res = cloudinary.api.resources(
        type="upload", prefix="story_books/", resource_type="raw", max_results=500
    ).get("resources", [])

    cover_map = {}
    try:
        img_res = cloudinary.api.resources(
            type="upload", prefix="story_covers/", resource_type="image", max_results=500
        ).get("resources", [])
        for r in img_res:
            cover_map[unquote(r["public_id"].replace("story_covers/", ""))] = r["secure_url"]
    except Exception:
        pass

    ch_cover_map = {}
    try:
        cc_res = cloudinary.api.resources(
            type="upload", prefix="story_chapter_covers/", resource_type="image", max_results=500
        ).get("resources", [])
        for r in cc_res:
            ch_cover_map[unquote(r["public_id"].replace("story_chapter_covers/", ""))] = r["secure_url"]
    except Exception:
        pass

    return raw_res, cover_map, ch_cover_map


@st.cache_data(ttl=300)
def fetch_docx_bytes(url):
    resp = requests.get(url, timeout=10)
    resp.raise_for_status()
    return resp.content


def parse_docx_bytes(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    chapters = []
    current_chapter = {"title": "", "music_url": "", "content": []}

    chapter_regex = re.compile(r"^第\s*[0-9一二三四五六七八九十百千]+\s*章")

    for p in doc.paragraphs:
        text = p.text.rstrip()
        style_name = ""
        if p.style and hasattr(p.style, "name") and p.style.name:
            style_name = p.style.name.lower()

        clean_text = text.strip()

        is_heading = (
            ("heading" in style_name or "標題" in style_name) and clean_text != ""
        ) or bool(chapter_regex.match(clean_text))

        if is_heading:
            if current_chapter["title"] or current_chapter["content"]:
                chapters.append(current_chapter)
            current_chapter = {
                "title": clean_text,
                "music_url": "",
                "content": [],
            }
        elif clean_text.startswith("http://") or clean_text.startswith("https://"):
            current_chapter["music_url"] = clean_text
        else:
            if not current_chapter["title"]:
                current_chapter["title"] = "前言/序章"
            current_chapter["content"].append(text)

    if current_chapter["title"] or current_chapter["content"]:
        chapters.append(current_chapter)

    if not chapters:
        chapters.append({
            "title": "全一冊",
            "music_url": "",
            "content": [p.text for p in doc.paragraphs]
        })

    return chapters


# 7. 音樂播放器輔助函式
def render_music_player(music_url, is_autoplay, is_loop):
    if music_url:
        if "youtube.com" in music_url or "youtu.be" in music_url:
            vid_match = re.search(r"(?:v=|be/|embed/)([\w-]+)", music_url)
            time_match = re.search(r"[?&](?:t|start)=(\d+)s?", music_url)

            if vid_match:
                vid = vid_match.group(1)
                embed_url = f"https://www.youtube.com/embed/{vid}"
                params = []
                if time_match:
                    params.append(f"start={time_match.group(1)}")
                if is_autoplay:
                    params.append("autoplay=1")
                if is_loop:
                    params.append(f"loop=1&playlist={vid}")

                if params:
                    embed_url += "?" + "&".join(params)

                st.sidebar.markdown(
                    f'<iframe width="100%" height="200" src="{embed_url}" frameborder="0" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen></iframe>',
                    unsafe_allow_html=True
                )
            else:
                st.sidebar.video(music_url, autoplay=is_autoplay, loop=is_loop)
        else:
            st.sidebar.audio(music_url, autoplay=is_autoplay, loop=is_loop)
    else:
        st.sidebar.caption("🎵 本章節未設定背景音樂")


has_cloud = init_cloudinary()

st.title("📚 雲端沉浸式故事音樂書櫃")

# 8. 主邏輯區域
if has_cloud:
    try:
        resources, cover_map, chapter_cover_map = fetch_cloudinary_catalog()

        book_options_map = {}
        for r in resources:
            b_title = unquote(r["public_id"].replace("story_books/", ""))
            book_options_map[b_title] = r["public_id"]

        # ---------------- 模式 A：總書櫃頁面 ----------------
        if not st.session_state.selected_book_id:
            st.sidebar.divider()
            st.sidebar.header("📤 故事入庫與覆蓋管理")

            upload_mode = st.sidebar.radio(
                "選擇操作模式：",
                ["✨ 全新上傳新書", "🔄 覆蓋現有書籍"],
                key="upload_mode_radio"
            )

            new_file = st.sidebar.file_uploader(
                "選擇 Word 故事檔 (.docx)", type=["docx"], key="main_file_uploader"
            )
            cover_file = st.sidebar.file_uploader(
                "選擇書籍封面圖片 (選填)", type=["png", "jpg", "jpeg"], key="main_cover_uploader"
            )

            if upload_mode == "✨ 全新上傳新書":
                if new_file:
                    if st.sidebar.button("💾 確認存入雲端書櫃", use_container_width=True):
                        with st.spinner("正在上傳故事與封面..."):
                            try:
                                safe_filename = quote(new_file.name)
                                cloudinary.uploader.upload(
                                    new_file,
                                    resource_type="raw",
                                    public_id=f"story_books/{safe_filename}",
                                    overwrite=True,
                                )
                                if cover_file:
                                    cloudinary.uploader.upload(
                                        cover_file,
                                        resource_type="image",
                                        public_id=f"story_covers/{safe_filename}",
                                        overwrite=True,
                                    )
                                st.sidebar.success(f"《{new_file.name}》與封面已成功收錄！")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.sidebar.error(f"上傳失敗：{e}")
            else:
                all_book_names = list(book_options_map.keys())
                if all_book_names:
                    target_to_overwrite = st.sidebar.selectbox(
                        "選擇要被覆蓋的書籍：", all_book_names, key="target_book_select"
                    )
                    confirm_overwrite = st.sidebar.checkbox(
                        "⚠️ 我確定要刪除原檔案並以新檔案覆蓋（更新上傳日期）", key="confirm_overwrite_box"
                    )

                    if new_file and confirm_overwrite:
                        if st.sidebar.button("💾 確認執行覆蓋", use_container_width=True):
                            with st.spinner("正在覆蓋雲端檔案與更新時間..."):
                                try:
                                    safe_filename = quote(target_to_overwrite)
                                    try:
                                        cloudinary.uploader.destroy(f"story_books/{safe_filename}", resource_type="raw", invalidate=True)
                                    except Exception:
                                        pass

                                    cloudinary.uploader.upload(
                                        new_file,
                                        resource_type="raw",
                                        public_id=f"story_books/{safe_filename}",
                                        overwrite=True,
                                        invalidate=True
                                    )
                                    if cover_file:
                                        cloudinary.uploader.upload(
                                            cover_file,
                                            resource_type="image",
                                            public_id=f"story_covers/{safe_filename}",
                                            overwrite=True,
                                        )
                                    st.sidebar.success(f"《{target_to_overwrite}》已成功覆蓋並更新上傳日期！")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.sidebar.error(f"覆蓋失敗：{e}")
                    elif new_file and not confirm_overwrite:
                        st.sidebar.warning("請先勾選上方「確認覆蓋」選項才能進行替換。")
                else:
                    st.sidebar.info("目前書櫃中尚無任何書籍可供覆蓋。")

            st.sidebar.divider()
            st.sidebar.header("🖼️ 各章節插圖管理")
            all_book_names_list = list(book_options_map.keys())
            if all_book_names_list:
                selected_lib_book = st.sidebar.selectbox("選擇要配圖的書籍：", all_book_names_list, key="lib_ch_book_select")
                try:
                    lib_book_res = next(r for r in resources if unquote(r["public_id"].replace("story_books/", "")) == selected_lib_book)
                    lib_bytes = fetch_docx_bytes(lib_book_res["secure_url"])
                    lib_chapters = parse_docx_bytes(lib_bytes)
                    lib_ch_titles = [c["title"] for c in lib_chapters]
                    selected_lib_ch = st.sidebar.selectbox("選擇要配圖的章節：", lib_ch_titles, key="lib_ch_select")

                    lib_ch_file = st.sidebar.file_uploader("選擇章節插圖", type=["png", "jpg", "jpeg"], key="lib_ch_file_up")
                    if st.sidebar.button("💾 上傳該章節插圖", use_container_width=True, key="lib_ch_btn"):
                        if lib_ch_file:
                            with st.spinner("上傳中..."):
                                target_ch_key = f"{selected_lib_book}_{selected_lib_ch}"
                                cloudinary.uploader.upload(
                                    lib_ch_file,
                                    resource_type="image",
                                    public_id=f"story_chapter_covers/{quote(target_ch_key)}",
                                    overwrite=True,
                                )
                                st.sidebar.success(f"《{selected_lib_book} - {selected_lib_ch}》插圖上傳成功！")
                                st.cache_data.clear()
                                st.rerun()
                        else:
                            st.sidebar.warning("請先選擇圖片檔案")
                except Exception:
                    st.sidebar.caption("載入章節失敗")

            if resources:
                st.header("📖 您的故事書櫃")
                cols = st.columns(3)

                for idx, res in enumerate(resources):
                    col = cols[idx % 3]
                    public_id = res["public_id"]
                    book_title = unquote(public_id.replace("story_books/", ""))
                    created_at_str = res.get("created_at", "")
                    date_display = (
                        datetime.strptime(
                            created_at_str, "%Y-%m-%dT%H:%M:%SZ"
                        ).strftime("%Y-%m-%d %H:%M")
                        if created_at_str
                        else "未知時間"
                    )

                    file_bytes_size = res.get("bytes", 0)
                    size_display = format_file_size(file_bytes_size)
                    cover_url = cover_map.get(book_title)

                    with col:
                        with st.container(border=True):
                            if cover_url:
                                st.image(cover_url, use_container_width=True)
                            else:
                                st.caption("📷 尚無封面圖片")

                            st.subheader(f"📘 {book_title}")
                            st.caption(f"📅 上傳：{date_display} ｜ 📦 大小：{size_display}")

                            b1, b2 = st.columns([2, 1])
                            with b1:
                                if st.button("📖 點擊閱讀", key=f"read_{public_id}", use_container_width=True):
                                    st.session_state.selected_book_id = public_id
                                    st.session_state.ch_index = 0
                                    st.session_state.reading_pct = 0
                                    st.rerun()
                            with b2:
                                if st.button("🗑️ 刪除", key=f"del_{public_id}", use_container_width=True):
                                    cloudinary.uploader.destroy(public_id, resource_type="raw")
                                    try:
                                        cloudinary.uploader.destroy(
                                            f"story_covers/{quote(book_title)}",
                                            resource_type="image",
                                        )
                                    except Exception:
                                        pass
                                    st.toast(f"已刪除《{book_title}》")
                                    st.cache_data.clear()
                                    st.rerun()

                            with st.popover("🖼️ 編輯/補傳封面", use_container_width=True):
                                edit_cover_file = st.file_uploader(
                                    "選擇封面圖片",
                                    type=["png", "jpg", "jpeg"],
                                    key=f"cov_up_{public_id}",
                                )
                                c_col1, c_col2 = st.columns(2)
                                with c_col1:
                                    if st.button("💾 覆蓋/補傳", key=f"btn_up_{public_id}", use_container_width=True):
                                        if edit_cover_file:
                                            with st.spinner("正在上傳封面..."):
                                                safe_filename = quote(book_title)
                                                cloudinary.uploader.upload(
                                                    edit_cover_file,
                                                    resource_type="image",
                                                    public_id=f"story_covers/{safe_filename}",
                                                    overwrite=True,
                                                )
                                                st.toast("封面已更新！")
                                                st.cache_data.clear()
                                                st.rerun()
                                        else:
                                            st.warning("請先選擇圖片檔案")
                                with c_col2:
                                    if st.button("🗑️ 刪除封面", key=f"btn_del_cov_{public_id}", use_container_width=True):
                                        try:
                                            safe_filename = quote(book_title)
                                            cloudinary.uploader.destroy(
                                                f"story_covers/{safe_filename}",
                                                resource_type="image",
                                            )
                                            st.toast("封面已刪除！")
                                            st.cache_data.clear()
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"刪除失敗：{e}")

        # ---------------- 模式 B：故事閱讀器頁面 ----------------
        else:
            selected_res = next(
                (r for r in resources if r["public_id"] == st.session_state.selected_book_id),
                None,
            )
            if selected_res:
                book_url = selected_res["secure_url"]
                book_name = unquote(selected_res["public_id"].replace("story_books/", ""))

                file_bytes = fetch_docx_bytes(book_url)
                chapters = parse_docx_bytes(file_bytes)

                st.session_state.max_chapters = len(chapters)
                st.session_state.chapter_titles = [ch["title"] for ch in chapters]

                if st.session_state.ch_index >= st.session_state.max_chapters:
                    st.session_state.ch_index = 0

                current_ch = chapters[st.session_state.ch_index]
                content_lines = current_ch["content"]
                total_lines = len(content_lines)

                # 側邊欄區域
                st.sidebar.divider()

                # 1. 本章插圖
                st.sidebar.caption(f"📖 本章插圖：{current_ch['title']}")
                ch_key = f"{book_name}_{current_ch['title']}"
                ch_cover_url = chapter_cover_map.get(ch_key)
                if ch_cover_url:
                    st.sidebar.image(ch_cover_url, use_container_width=True)
                else:
                    st.sidebar.caption("📷 本章尚無插圖")

                with st.sidebar.popover("🖼️ 上傳/更換本章插圖", use_container_width=True):
                    ch_img_file = st.file_uploader("選擇本章插圖", type=["png", "jpg", "jpeg"], key=f"ch_up_{ch_key}")
                    if st.button("💾 確認上傳本章插圖", key=f"btn_ch_up_{ch_key}", use_container_width=True):
                        if ch_img_file:
                            with st.spinner("正在上傳本章插圖..."):
                                cloudinary.uploader.upload(
                                    ch_img_file,
                                    resource_type="image",
                                    public_id=f"story_chapter_covers/{quote(ch_key)}",
                                    overwrite=True,
                                )
                                st.toast("本章插圖已更新！")
                                st.cache_data.clear()
                                st.rerun()
                        else:
                            st.warning("請先選擇圖片檔案")

                # 2. 書籍封面
                st.sidebar.divider()
                st.sidebar.caption(f"📘 書籍封面：{book_name}")
                reading_cover_url = cover_map.get(book_name)
                if reading_cover_url:
                    st.sidebar.image(reading_cover_url, use_container_width=True)
                else:
                    st.sidebar.caption("📷 本書尚無封面圖片")

                # 3. 下載本書 Word 檔按鈕
                st.sidebar.divider()
                st.sidebar.header("📥 檔案下載")
                st.sidebar.download_button(
                    label="📥 下載本書 Word 檔",
                    data=file_bytes,
                    file_name=book_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="sidebar_download_docx"
                )

                # 4. 音樂盒與循環播放控制
                st.sidebar.divider()
                st.sidebar.header("🎵 音樂盒")

                with st.sidebar.container(border=True):
                    st.session_state.auto_play = st.toggle(
                        "▶️ 切換章節自動播放音樂",
                        value=st.session_state.auto_play,
                        key="auto_play_toggle"
                    )
                    st.session_state.loop_play = st.toggle(
                        "🔁 音樂循環播放",
                        value=st.session_state.loop_play,
                        key="loop_play_toggle"
                    )

                render_music_player(
                    current_ch["music_url"], st.session_state.auto_play, st.session_state.loop_play
                )

                st.sidebar.divider()

                st.sidebar.header("📌 章節切換與進度")

                pct_value = st.sidebar.slider(
                    "🎯 快轉跳轉：",
                    min_value=0,
                    max_value=100,
                    value=st.session_state.reading_pct,
                    step=5,
                    format="%d%%",
                    key="sb_slider_pct",
                    on_change=on_slider_change_cb,
                )

                target_line_idx = int(total_lines * (pct_value / 100.0))
                if target_line_idx >= total_lines:
                    target_line_idx = max(0, total_lines - 1)

                with st.sidebar.popover(f"跳轉章節：{current_ch['title']}", use_container_width=True):
                    st.radio(
                        "請選擇要閱讀的章節：",
                        st.session_state.chapter_titles,
                        index=st.session_state.ch_index,
                        key="sb_popover_radio",
                        on_change=on_radio_change_cb,
                    )

                nav_c1, nav_c2 = st.sidebar.columns(2)
                with nav_c1:
                    st.button(
                        "⬅️ 上一章",
                        disabled=(st.session_state.ch_index <= 0),
                        use_container_width=True,
                        key="side_prev",
                        on_click=prev_chapter_cb,
                    )
                with nav_c2:
                    st.button(
                        "下一章 ➡️",
                        disabled=(st.session_state.ch_index >= st.session_state.max_chapters - 1),
                        use_container_width=True,
                        key="side_next",
                        on_click=next_chapter_cb,
                    )

                st.sidebar.divider()

                # 控制面板
                st.sidebar.header("🎛️ 控制面板")

                all_book_names = list(book_options_map.keys())
                if all_book_names:
                    current_book_idx = (
                        all_book_names.index(book_name)
                        if book_name in all_book_names
                        else 0
                    )
                    selected_target_book = st.sidebar.selectbox(
                        "📚 快速切換書籍：",
                        all_book_names,
                        index=current_book_idx,
                        key="sidebar_book_switch",
                    )
                    if book_options_map.get(selected_target_book) != st.session_state.selected_book_id:
                        st.session_state.selected_book_id = book_options_map[selected_target_book]
                        st.session_state.ch_index = 0
                        st.session_state.reading_pct = 0
                        st.rerun()

                if st.sidebar.button("📚 返回圖書總書櫃", use_container_width=True):
                    st.session_state.selected_book_id = None
                    st.session_state.reading_pct = 0
                    st.rerun()

                # ---------------- 文章閱讀主區域 ----------------
                st.header(f"《{book_name}》")
                st.subheader(current_ch["title"])
                st.divider()

                # 頂部控制列：包含「上一章」、「快速選章 Popover」、「下一章」
                top_c1, top_c2, top_c3 = st.columns([1, 2, 1])
                with top_c1:
                    st.button(
                        "⬅️ 上一章",
                        disabled=(st.session_state.ch_index <= 0),
                        use_container_width=True,
                        key="top_prev",
                        on_click=prev_chapter_cb,
                    )
                with top_c2:
                    # 補回頂部的文章內跳章選單
                    with st.popover(f"📑 快速跳章 ({st.session_state.ch_index + 1}/{st.session_state.max_chapters})", use_container_width=True):
                        st.radio(
                            "請選擇要閱讀的章節：",
                            st.session_state.chapter_titles,
                            index=st.session_state.ch_index,
                            key="top_popover_radio",
                            on_change=on_radio_change_cb,
                        )
                with top_c3:
                    st.button(
                        "下一章 ➡️",
                        disabled=(st.session_state.ch_index >= st.session_state.max_chapters - 1),
                        use_container_width=True,
                        key="top_next",
                        on_click=next_chapter_cb,
                    )

                st.divider()

                # 逐段渲染內文：使用 HTML <p class="story-paragraph"> 強制分開段落
                for idx, line in enumerate(content_lines):
                    anchor_html = f'<div id="line-anchor-{idx}"></div>'
                    st.markdown(anchor_html, unsafe_allow_html=True)
                    
                    line_str = str(line)
                    if line_str.strip() == "":
                        # 空行增加適當距離
                        st.markdown('<div style="height: 1rem;"></div>', unsafe_allow_html=True)
                    else:
                        # 非空行獨立封裝為 HTML 段落
                        st.markdown(f'<p class="story-paragraph">{line_str}</p>', unsafe_allow_html=True)

                st.divider()

                bot_c1, bot_c2 = st.columns(2)
                with bot_c1:
                    st.button(
                        "⬅️ 上一章",
                        disabled=(st.session_state.ch_index <= 0),
                        use_container_width=True,
                        key="bot_prev",
                        on_click=prev_chapter_cb,
                    )
                with bot_c2:
                    st.button(
                        "下一章 ➡️",
                        disabled=(st.session_state.ch_index >= st.session_state.max_chapters - 1),
                        use_container_width=True,
                        key="bot_next",
                        on_click=next_chapter_cb,
                    )

    except Exception as e:
        st.error(f"連線至雲端書櫃時發生錯誤：{e}")
else:
    st.warning("⚠️ 請先在 Streamlit Cloud 設定 Cloudinary API 金鑰。")
